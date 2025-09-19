"""
Document parsing module for healthcare requirements.
Supports PDF, Word, XML, Markdown, and other formats.
"""

import os
import hashlib
import tempfile
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import logging
from datetime import datetime

# Document parsing libraries
import pypdf
from docx import Document
import xmltodict
import markdown
from bs4 import BeautifulSoup
import pandas as pd
import openpyxl

from models import (
    DocumentType, DocumentMetadata, GraphState, ProcessingStatus, WorkflowStep
)
from config import settings, HealthcareDomainConfig

logger = logging.getLogger(__name__)


class DocumentParser:
    """Main document parser for healthcare requirements."""
    
    def __init__(self):
        self.supported_types = {
            DocumentType.PDF: self._parse_pdf,
            DocumentType.WORD: self._parse_word,
            DocumentType.XML: self._parse_xml,
            DocumentType.MARKDOWN: self._parse_markdown,
            DocumentType.TEXT: self._parse_text,
            DocumentType.EXCEL: self._parse_excel
        }
    
    def parse_documents(self, state: GraphState) -> GraphState:
        """
        Parse multiple documents and extract text content.
        
        Args:
            state: Current workflow state
            
        Returns:
            Updated state with parsed content
        """
        logger.info(f"Starting document parsing for {len(state.input_documents)} documents")
        
        # Create workflow step
        step = WorkflowStep(
            step_name="document_parser",
            status=ProcessingStatus.IN_PROGRESS,
            start_time=datetime.now()
        )
        state.workflow_steps.append(step)
        state.current_step = "document_parser"
        
        try:
            parsed_documents = []
            metadata_list = []
            raw_content_list = []
            
            for doc_data in state.input_documents:
                try:
                    # Parse individual document
                    parsed_doc = self._parse_single_document(doc_data)
                    if parsed_doc:
                        parsed_documents.append(parsed_doc)
                        metadata_list.append(parsed_doc['metadata'])
                        raw_content_list.append(parsed_doc['content'])
                        
                except Exception as e:
                    error_msg = f"Failed to parse document {doc_data.get('filename', 'unknown')}: {str(e)}"
                    logger.error(error_msg)
                    state.error_log.append(error_msg)
            
            # Update state
            state.document_metadata = metadata_list
            state.raw_text_content = raw_content_list
            state.overall_status = ProcessingStatus.COMPLETED if parsed_documents else ProcessingStatus.FAILED
            
            # Update workflow step
            step.status = ProcessingStatus.COMPLETED
            step.end_time = datetime.now()
            step.duration_seconds = (step.end_time - step.start_time).total_seconds()
            step.output_data = {
                "parsed_documents_count": len(parsed_documents),
                "total_content_length": sum(len(content) for content in raw_content_list)
            }
            
            logger.info(f"Successfully parsed {len(parsed_documents)} documents")
            
        except Exception as e:
            error_msg = f"Document parsing failed: {str(e)}"
            logger.error(error_msg)
            state.error_log.append(error_msg)
            state.overall_status = ProcessingStatus.FAILED
            
            step.status = ProcessingStatus.FAILED
            step.end_time = datetime.now()
            step.error_message = error_msg
        
        return state
    
    def _parse_single_document(self, doc_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Parse a single document and return content with metadata."""
        try:
            file_path = doc_data.get('file_path')
            filename = doc_data.get('filename', Path(file_path).name if file_path else 'unknown')
            content = doc_data.get('content')
            
            if not content and not file_path:
                raise ValueError("Either file_path or content must be provided")
            
            # Determine document type
            doc_type = self._detect_document_type(filename, content)
            
            # Create metadata
            metadata = self._create_metadata(filename, doc_type, content, file_path)
            
            # Parse content based on type
            if content:
                parsed_content = self._parse_content_by_type(content, doc_type)
            else:
                parsed_content = self._parse_file_by_type(file_path, doc_type)
            
            # Clean and preprocess content
            cleaned_content = self._clean_content(parsed_content)
            
            return {
                'metadata': metadata,
                'content': cleaned_content,
                'original_type': doc_type
            }
            
        except Exception as e:
            logger.error(f"Failed to parse document {filename}: {str(e)}")
            return None
    
    def _detect_document_type(self, filename: str, content: Optional[str] = None) -> DocumentType:
        """Detect document type based on filename and content."""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return DocumentType.PDF
        elif filename_lower.endswith(('.doc', '.docx')):
            return DocumentType.WORD
        elif filename_lower.endswith(('.xml', '.xsd')):
            return DocumentType.XML
        elif filename_lower.endswith(('.md', '.markdown')):
            return DocumentType.MARKDOWN
        elif filename_lower.endswith(('.xlsx', '.xls')):
            return DocumentType.EXCEL
        elif filename_lower.endswith('.txt'):
            return DocumentType.TEXT
        else:
            # Try to detect from content
            if content:
                if content.strip().startswith('<?xml'):
                    return DocumentType.XML
                elif content.strip().startswith('#'):
                    return DocumentType.MARKDOWN
            
            return DocumentType.TEXT  # Default fallback
    
    def _create_metadata(self, filename: str, doc_type: DocumentType, 
                        content: Optional[str], file_path: Optional[str]) -> DocumentMetadata:
        """Create document metadata."""
        file_size = 0
        checksum = ""
        
        if content:
            file_size = len(content.encode('utf-8'))
            checksum = hashlib.md5(content.encode('utf-8')).hexdigest()
        elif file_path and os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            with open(file_path, 'rb') as f:
                checksum = hashlib.md5(f.read()).hexdigest()
        
        return DocumentMetadata(
            filename=filename,
            document_type=doc_type,
            file_size=file_size,
            upload_timestamp=datetime.now(),
            checksum=checksum
        )
    
    def _parse_content_by_type(self, content: str, doc_type: DocumentType) -> str:
        """Parse content based on document type."""
        parser_func = self.supported_types.get(doc_type)
        if parser_func:
            return parser_func(content)
        else:
            return content
    
    def _parse_file_by_type(self, file_path: str, doc_type: DocumentType) -> str:
        """Parse file based on document type."""
        parser_func = self.supported_types.get(doc_type)
        if parser_func:
            with open(file_path, 'rb') as f:
                return parser_func(f.read())
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    
    def _parse_pdf(self, content: Union[str, bytes]) -> str:
        """Parse PDF content."""
        try:
            if isinstance(content, str):
                # If content is base64 encoded or similar, decode it
                import base64
                content = base64.b64decode(content)
            
            # Create temporary file for PDF parsing
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                # Parse PDF
                with open(temp_file.name, 'rb') as pdf_file:
                    pdf_reader = pypdf.PdfReader(pdf_file)
                    text_content = []
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                        except Exception as e:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    
                    return "\n\n".join(text_content)
        
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            return str(content) if isinstance(content, str) else ""
    
    def _parse_word(self, content: Union[str, bytes]) -> str:
        """Parse Word document content."""
        try:
            if isinstance(content, str):
                # If content is base64 encoded, decode it
                import base64
                content = base64.b64decode(content)
            
            # Create temporary file for Word parsing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                # Parse Word document
                doc = Document(temp_file.name)
                text_content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                return "\n".join(text_content)
        
        except Exception as e:
            logger.error(f"Word document parsing failed: {str(e)}")
            return str(content) if isinstance(content, str) else ""
    
    def _parse_xml(self, content: Union[str, bytes]) -> str:
        """Parse XML content."""
        try:
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            # Parse XML to dictionary
            xml_dict = xmltodict.parse(content)
            
            # Convert to readable text format
            text_content = self._xml_dict_to_text(xml_dict)
            return text_content
        
        except Exception as e:
            logger.error(f"XML parsing failed: {str(e)}")
            return str(content) if isinstance(content, str) else ""
    
    def _xml_dict_to_text(self, xml_dict: Dict[str, Any], indent: int = 0) -> str:
        """Convert XML dictionary to readable text."""
        text_parts = []
        indent_str = "  " * indent
        
        for key, value in xml_dict.items():
            if isinstance(value, dict):
                text_parts.append(f"{indent_str}{key}:")
                text_parts.append(self._xml_dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                for i, item in enumerate(value):
                    text_parts.append(f"{indent_str}{key}[{i}]:")
                    if isinstance(item, dict):
                        text_parts.append(self._xml_dict_to_text(item, indent + 1))
                    else:
                        text_parts.append(f"{indent_str}  {item}")
            else:
                text_parts.append(f"{indent_str}{key}: {value}")
        
        return "\n".join(text_parts)
    
    def _parse_markdown(self, content: Union[str, bytes]) -> str:
        """Parse Markdown content."""
        try:
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            # Convert markdown to HTML
            html = markdown.markdown(content)
            
            # Extract text from HTML
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        
        except Exception as e:
            logger.error(f"Markdown parsing failed: {str(e)}")
            return str(content) if isinstance(content, str) else ""
    
    def _parse_text(self, content: Union[str, bytes]) -> str:
        """Parse plain text content."""
        try:
            if isinstance(content, bytes):
                return content.decode('utf-8')
            return content
        except Exception as e:
            logger.error(f"Text parsing failed: {str(e)}")
            return str(content) if isinstance(content, str) else ""
    
    def _parse_excel(self, content: Union[str, bytes]) -> str:
        """Parse Excel content."""
        try:
            if isinstance(content, str):
                # If content is base64 encoded, decode it
                import base64
                content = base64.b64decode(content)
            
            # Create temporary file for Excel parsing
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                # Parse Excel file
                df = pd.read_excel(temp_file.name, sheet_name=None)
                text_content = []
                
                for sheet_name, sheet_df in df.items():
                    text_content.append(f"--- Sheet: {sheet_name} ---")
                    text_content.append(sheet_df.to_string(index=False))
                    text_content.append("")
                
                return "\n".join(text_content)
        
        except Exception as e:
            logger.error(f"Excel parsing failed: {str(e)}")
            return str(content) if isinstance(content, str) else ""
    
    def _clean_content(self, content: str) -> str:
        """Clean and preprocess extracted content."""
        if not content:
            return ""
        
        # Remove excessive whitespace
        content = ' '.join(content.split())
        
        # Remove special characters that might interfere with processing
        import re
        content = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\\]', ' ', content)
        
        # Remove multiple spaces
        content = re.sub(r'\s+', ' ', content)
        
        # Remove empty lines
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        
        return '\n'.join(lines)
    
    def extract_requirements_sections(self, content: str) -> Dict[str, str]:
        """Extract different sections from healthcare requirements documents."""
        sections = {
            'functional_requirements': '',
            'non_functional_requirements': '',
            'compliance_requirements': '',
            'security_requirements': '',
            'performance_requirements': '',
            'usability_requirements': ''
        }
        
        # Common section headers in healthcare documents
        section_patterns = {
            'functional_requirements': [
                r'functional\s+requirements?',
                r'business\s+requirements?',
                r'system\s+functionality',
                r'features?\s+and\s+capabilities'
            ],
            'non_functional_requirements': [
                r'non[- ]?functional\s+requirements?',
                r'technical\s+requirements?',
                r'system\s+requirements?'
            ],
            'compliance_requirements': [
                r'compliance\s+requirements?',
                r'regulatory\s+requirements?',
                r'standards?\s+compliance',
                r'FDA\s+requirements?',
                r'IEC\s+62304',
                r'ISO\s+\d+'
            ],
            'security_requirements': [
                r'security\s+requirements?',
                r'data\s+protection',
                r'privacy\s+requirements?',
                r'HIPAA\s+compliance',
                r'access\s+control'
            ],
            'performance_requirements': [
                r'performance\s+requirements?',
                r'scalability\s+requirements?',
                r'response\s+time',
                r'throughput\s+requirements?'
            ],
            'usability_requirements': [
                r'usability\s+requirements?',
                r'user\s+interface\s+requirements?',
                r'user\s+experience',
                r'ergonomics'
            ]
        }
        
        import re
        content_lower = content.lower()
        
        for section_name, patterns in section_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, content_lower)
                if match:
                    # Extract content after the section header
                    start_pos = match.end()
                    # Find next section or end of document
                    next_section_pos = len(content)
                    for other_patterns in section_patterns.values():
                        for other_pattern in other_patterns:
                            if other_pattern != pattern:
                                other_match = re.search(other_pattern, content_lower[start_pos:])
                                if other_match:
                                    next_section_pos = min(next_section_pos, start_pos + other_match.start())
                    
                    sections[section_name] = content[start_pos:next_section_pos].strip()
                    break
        
        return sections


# Export the main parser class
__all__ = ["DocumentParser"]
